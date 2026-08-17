"""
File management service for organizing processed news.
Creates structured folder hierarchy and saves Word documents and images.
"""
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict
from docx import Document

from config import OUTPUT_DIR


class FileManager:
    """Service for managing news file storage."""
    
    def __init__(self):
        self.output_dir = Path(OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _sanitize_filename(self, text: str, max_length: int = 50) -> str:
        """
        Sanitize text to be used as filename.
        
        Args:
            text: Original text
            max_length: Maximum length of filename
            
        Returns:
            Sanitized filename
        """
        # Remove special characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '', text)
        # Replace spaces with underscores
        sanitized = sanitized.replace(' ', '_')
        # Remove consecutive underscores
        sanitized = re.sub(r'_+', '_', sanitized)
        # Trim to max length
        sanitized = sanitized[:max_length]
        # Remove trailing underscores
        sanitized = sanitized.strip('_')
        
        return sanitized if sanitized else 'untitled'
    
    def create_news_folder(
        self,
        parse_date: datetime,
        title: str
    ) -> Path:
        """
        Create folder structure for a news item.
        
        Structure: output/YYYY-MM-DD/Title_Text/
        
        Args:
            parse_date: Date when news was parsed
            title: News title
            
        Returns:
            Path to the created folder
        """
        # Format date as YYYY-MM-DD
        date_str = parse_date.strftime('%Y-%m-%d')
        
        # Create date folder
        date_folder = self.output_dir / date_str
        date_folder.mkdir(parents=True, exist_ok=True)
        
        # Sanitize title for folder name
        safe_title = self._sanitize_filename(title, max_length=80)
        
        # Create title folder
        news_folder = date_folder / safe_title
        
        # Handle duplicate folder names
        counter = 1
        original_folder = news_folder
        while news_folder.exists():
            news_folder = date_folder / f"{safe_title}_{counter}"
            counter += 1
        
        news_folder.mkdir(parents=True, exist_ok=True)
        
        return news_folder
    
    def save_word_document(
        self,
        folder_path: Path,
        title: str,
        content: str,
        metadata: Optional[Dict] = None
    ) -> Optional[str]:
        """
        Save news article as Word document.
        
        Args:
            folder_path: Folder to save document in
            title: Article title
            content: Article content
            metadata: Optional metadata (source, date, link, etc.)
            
        Returns:
            Full path to saved file or None if failed
        """
        try:
            # Create document
            doc = Document()
            
            # Add title
            doc.add_heading(title, level=1)
            
            # Add metadata if provided
            if metadata:
                doc.add_paragraph()
                if 'source' in metadata:
                    doc.add_paragraph(f"Источник: {metadata['source']}", style='Intense Quote')
                if 'date' in metadata:
                    doc.add_paragraph(f"Дата публикации: {metadata['date']}", style='Intense Quote')
                if 'link' in metadata:
                    doc.add_paragraph(f"Ссылка: {metadata['link']}", style='Intense Quote')
                doc.add_paragraph()
            
            # Add content
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Save document
            doc_path = folder_path / 'article.docx'
            doc.save(doc_path)
            
            return str(doc_path)
            
        except Exception as e:
            print(f"Error saving Word document: {str(e)}")
            return None
    
    def save_image(
        self,
        folder_path: Path,
        image_bytes: bytes,
        filename: str = 'image'
    ) -> Optional[str]:
        """
        Save image file.
        
        Args:
            folder_path: Folder to save image in
            image_bytes: Image data
            filename: Base filename (without extension)
            
        Returns:
            Full path to saved file or None if failed
        """
        try:
            # Save as JPG
            image_path = folder_path / f'{filename}.jpg'
            
            with open(image_path, 'wb') as f:
                f.write(image_bytes)
            
            return str(image_path)
            
        except Exception as e:
            print(f"Error saving image: {str(e)}")
            return None
    
    def save_complete_news(
        self,
        parse_date: datetime,
        title: str,
        content: str,
        image_bytes: Optional[bytes],
        metadata: Optional[Dict] = None
    ) -> Dict[str, Optional[str]]:
        """
        Save complete news item (folder + Word doc + image).
        
        Args:
            parse_date: Date when news was parsed
            title: Article title
            content: Article content
            image_bytes: Generated image data
            metadata: Optional metadata
            
        Returns:
            Dictionary with paths to saved files
        """
        result = {
            'folder': None,
            'document': None,
            'image': None,
            'success': False
        }
        
        try:
            # Create folder structure
            folder = self.create_news_folder(parse_date, title)
            result['folder'] = str(folder)
            
            # Save Word document
            doc_path = self.save_word_document(folder, title, content, metadata)
            result['document'] = doc_path
            
            # Save image if provided
            if image_bytes:
                image_path = self.save_image(folder, image_bytes, 'image')
                result['image'] = image_path
            
            # Mark as successful if at least document was saved
            if doc_path:
                result['success'] = True
            
        except Exception as e:
            print(f"Error saving complete news: {str(e)}")
        
        return result
    
    def get_recent_folders(self, limit: int = 10) -> list:
        """
        Get list of recently created news folders.
        
        Args:
            limit: Maximum number of folders to return
            
        Returns:
            List of folder paths, newest first
        """
        folders = []
        
        # Iterate through date folders
        for date_folder in sorted(self.output_dir.iterdir(), reverse=True):
            if not date_folder.is_dir():
                continue
            
            # Iterate through news folders within date folder
            for news_folder in sorted(date_folder.iterdir(), reverse=True):
                if news_folder.is_dir():
                    folders.append(news_folder)
                    
                    if len(folders) >= limit:
                        return folders
        
        return folders
    
    def cleanup_old_news(self, days_to_keep: int = 30) -> int:
        """
        Remove news folders older than specified days.
        
        Args:
            days_to_keep: Number of days to keep news
            
        Returns:
            Number of folders removed
        """
        from datetime import timedelta
        
        cutoff_date = datetime.now() - timedelta(days=days_to_keep)
        removed_count = 0
        
        for date_folder in self.output_dir.iterdir():
            if not date_folder.is_dir():
                continue
            
            try:
                # Parse date from folder name
                folder_date = datetime.strptime(date_folder.name, '%Y-%m-%d')
                
                if folder_date < cutoff_date:
                    # Remove entire date folder
                    import shutil
                    shutil.rmtree(date_folder)
                    removed_count += 1
                    
            except ValueError:
                # Not a date folder, skip
                continue
        
        return removed_count


# Example usage
if __name__ == '__main__':
    def main():
        fm = FileManager()
        
        # Test saving a news item
        result = fm.save_complete_news(
            parse_date=datetime.now(),
            title="Тестовая новость о технологиях",
            content="Это пример содержимого новости. Здесь должен быть полный текст статьи.",
            image_bytes=None,  # No image for this test
            metadata={
                'source': 'Test Source',
                'date': datetime.now().isoformat(),
                'link': 'https://example.com/news/123'
            }
        )
        
        print(f"Folder: {result['folder']}")
        print(f"Document: {result['document']}")
        print(f"Image: {result['image']}")
        print(f"Success: {result['success']}")
    
    main()
